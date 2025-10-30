# services/ocr_service.py
import os
import logging
import easyocr
import fitz  # PyMuPDF
import docx
import pandas as pd
import re
import io
from typing import Dict, Any, Tuple, List, Optional
import tempfile
from PIL import Image
import pytesseract
import numpy as np
import magic
from datetime import datetime

logger = logging.getLogger(__name__)

class OCRService:
    """
    Service for extracting text and content from various document formats
    using OCR and other extraction techniques.
    """
    
    def __init__(self, use_gpu: bool = False):
        """
        Initialize the OCR service
        
        Args:
            use_gpu: Whether to use GPU acceleration for OCR (requires CUDA)
        """
        self.use_gpu = use_gpu
        self._initialize_ocr_engines()
        logger.info(f"OCR Service initialized with GPU support: {use_gpu}")
    
    def _initialize_ocr_engines(self):
        """Initialize OCR engines"""
        # Initialize EasyOCR reader for multiple languages
        try:
            self.reader = easyocr.Reader(['en'], gpu=self.use_gpu)
            logger.info("EasyOCR engine initialized successfully")
        except Exception as e:
            logger.error(f"Error initializing EasyOCR: {str(e)}")
            self.reader = None
        
        # Initialize Tesseract path if needed
        if os.environ.get('TESSERACT_PATH'):
            pytesseract.pytesseract.tesseract_cmd = os.environ.get('TESSERACT_PATH')
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a file based on its type
        
        Args:
            file_content: Binary content of the file
            filename: Name of the file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        # Detect file type
        mime_type = self._detect_mime_type(file_content)
        file_ext = os.path.splitext(filename.lower())[1]
        
        logger.info(f"Extracting text from file: {filename} (MIME: {mime_type})")
        
        # Extract based on file type
        if mime_type.startswith('image/'):
            text, metadata = self.extract_text_from_image(file_content)
            metadata['extraction_method'] = 'image_ocr'
        
        elif mime_type == 'application/pdf' or file_ext == '.pdf':
            text, metadata = self.extract_text_from_pdf(file_content)
            metadata['extraction_method'] = 'pdf_extraction'
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or file_ext == '.docx':
            text, metadata = self.extract_text_from_docx(file_content)
            metadata['extraction_method'] = 'docx_extraction'
        
        elif mime_type == 'text/markdown' or file_ext in ['.md', '.markdown']:
            text = file_content.decode('utf-8', errors='replace')
            metadata = {'extraction_method': 'text_extraction'}
        
        elif mime_type.startswith('text/') or file_ext in ['.txt', '.csv', '.json', '.xml', '.html', '.htm']:
            text = file_content.decode('utf-8', errors='replace')
            metadata = {'extraction_method': 'text_extraction'}
        
        elif file_ext in ['.py', '.js', '.java', '.c', '.cpp', '.cs', '.go', '.rb', '.php', '.swift']:
            text, metadata = self.extract_code_comments(file_content, file_ext)
            metadata['extraction_method'] = 'code_extraction'
        
        elif file_ext in ['.xlsx', '.xls']:
            text, metadata = self.extract_text_from_excel(file_content)
            metadata['extraction_method'] = 'excel_extraction'
        
        else:
            # Default to trying as text
            try:
                text = file_content.decode('utf-8', errors='replace')
                metadata = {'extraction_method': 'fallback_text_extraction'}
            except UnicodeDecodeError:
                text = f"[Could not extract text from {filename}]"
                metadata = {'extraction_method': 'failed', 'error': 'Unsupported format'}
        
        # Add common metadata
        metadata['filename'] = filename
        metadata['file_size'] = len(file_content)
        metadata['mime_type'] = mime_type
        metadata['extraction_date'] = datetime.utcnow().isoformat()
        
        return text, metadata
    
    def extract_text_from_image(self, image_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from an image using OCR
        
        Args:
            image_content: Binary content of the image
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {}
        
        try:
            # Convert bytes to image
            image = Image.open(io.BytesIO(image_content))
            
            # Add image metadata
            metadata['image_width'] = image.width
            metadata['image_height'] = image.height
            metadata['image_format'] = image.format
            metadata['image_mode'] = image.mode
            
            # Use EasyOCR if available
            if self.reader:
                # Convert PIL Image to numpy array
                image_np = np.array(image)
                
                # Run OCR
                results = self.reader.readtext(image_np)
                
                # Extract text
                text = ' '.join([result[1] for result in results])
                
                # Add confidence scores to metadata
                metadata['ocr_engine'] = 'easyocr'
                metadata['confidence_scores'] = [float(result[2]) for result in results]
                metadata['avg_confidence'] = sum(metadata['confidence_scores']) / len(metadata['confidence_scores']) if metadata['confidence_scores'] else 0
            
            # Fallback to Tesseract
            else:
                text = pytesseract.image_to_string(image)
                metadata['ocr_engine'] = 'tesseract'
            
            logger.info(f"Successfully extracted text from image: {len(text)} characters")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            return f"[OCR Error: {str(e)}]", {'error': str(e), 'ocr_engine': 'failed'}
    
    def extract_text_from_pdf(self, pdf_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from a PDF file
        
        Args:
            pdf_content: Binary content of the PDF
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {}
        full_text = []
        
        try:
            # Load PDF
            pdf_document = fitz.open(stream=pdf_content, filetype="pdf")
            
            # Extract document metadata
            pdf_metadata = pdf_document.metadata
            if pdf_metadata:
                metadata['title'] = pdf_metadata.get('title', '')
                metadata['author'] = pdf_metadata.get('author', '')
                metadata['subject'] = pdf_metadata.get('subject', '')
                metadata['creator'] = pdf_metadata.get('creator', '')
                metadata['producer'] = pdf_metadata.get('producer', '')
                metadata['creation_date'] = pdf_metadata.get('creationDate', '')
                metadata['modification_date'] = pdf_metadata.get('modDate', '')
            
            metadata['page_count'] = len(pdf_document)
            metadata['has_text'] = False
            
            # Extract text from each page
            for page_num, page in enumerate(pdf_document):
                text = page.get_text()
                
                # If page has no text, try OCR
                if not text.strip():
                    # Convert page to image
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    
                    # Use OCR to extract text from the image
                    page_text, _ = self.extract_text_from_image(img_data)
                    full_text.append(page_text)
                    
                    if not metadata.get('pages_requiring_ocr'):
                        metadata['pages_requiring_ocr'] = []
                    metadata['pages_requiring_ocr'].append(page_num + 1)
                else:
                    full_text.append(text)
                    metadata['has_text'] = True
            
            # Extract tables if available
            tables = self.extract_table_structure(pdf_content)
            if tables:
                metadata['tables_extracted'] = len(tables)
                # Add table text to full text
                for table in tables:
                    full_text.append(table)
            
            # Combine all text
            combined_text = "\n\n".join(full_text)
            
            logger.info(f"Successfully extracted text from PDF: {len(combined_text)} characters, {metadata['page_count']} pages")
            return combined_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            return f"[PDF Extraction Error: {str(e)}]", {'error': str(e)}
    
    def extract_text_from_docx(self, docx_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from a DOCX file
        
        Args:
            docx_content: Binary content of the DOCX file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {}
        
        try:
            # Create a temporary file to load the docx
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_content)
                temp_file_path = temp_file.name
            
            # Load the document
            doc = docx.Document(temp_file_path)
            
            # Extract document properties
            core_properties = doc.core_properties
            metadata['title'] = core_properties.title if hasattr(core_properties, 'title') else ''
            metadata['author'] = core_properties.author if hasattr(core_properties, 'author') else ''
            metadata['created'] = str(core_properties.created) if hasattr(core_properties, 'created') else ''
            metadata['modified'] = str(core_properties.modified) if hasattr(core_properties, 'modified') else ''
            metadata['last_modified_by'] = core_properties.last_modified_by if hasattr(core_properties, 'last_modified_by') else ''
            
            # Extract content
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                # Convert table to string representation
                table_str = "\n".join([" | ".join(row) for row in table_data])
                tables.append(table_str)
            
            # Combine all content
            full_text = "\n\n".join(paragraphs + tables)
            
            # Add additional metadata
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            # Clean up temp file
            os.unlink(temp_file_path)
            
            logger.info(f"Successfully extracted text from DOCX: {len(full_text)} characters")
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            return f"[DOCX Extraction Error: {str(e)}]", {'error': str(e)}
    
    def extract_text_from_excel(self, excel_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from Excel files
        
        Args:
            excel_content: Binary content of the Excel file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {}
        
        try:
            # Read Excel file
            excel_file = io.BytesIO(excel_content)
            dfs = pd.read_excel(excel_file, sheet_name=None)
            
            metadata['sheet_count'] = len(dfs)
            metadata['sheet_names'] = list(dfs.keys())
            
            # Process each sheet
            all_text = []
            for sheet_name, df in dfs.items():
                # Convert sheet to string
                sheet_text = f"--- Sheet: {sheet_name} ---\n"
                sheet_text += df.to_string(index=False)
                all_text.append(sheet_text)
            
            # Combine all sheets
            combined_text = "\n\n".join(all_text)
            
            logger.info(f"Successfully extracted text from Excel: {len(combined_text)} characters, {metadata['sheet_count']} sheets")
            return combined_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text from Excel: {str(e)}")
            return f"[Excel Extraction Error: {str(e)}]", {'error': str(e)}
    
    def extract_code_comments(self, code_content: bytes, file_ext: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract code and comments from source code files
        
        Args:
            code_content: Binary content of the code file
            file_ext: File extension to determine language
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        metadata = {}
        
        try:
            # Convert bytes to string
            code_text = code_content.decode('utf-8', errors='replace')
            
            # Determine language based on file extension
            language = file_ext.lstrip('.')
            metadata['language'] = language
            
            # Define comment patterns based on language
            single_line_comment = None
            multi_line_start = None
            multi_line_end = None
            
            if language in ['py', 'python']:
                single_line_comment = '#'
                multi_line_start = '"""'
                multi_line_end = '"""'
            elif language in ['js', 'javascript', 'ts', 'typescript', 'java', 'c', 'cpp', 'cs', 'go', 'php', 'swift']:
                single_line_comment = '//'
                multi_line_start = '/*'
                multi_line_end = '*/'
            elif language in ['rb', 'ruby']:
                single_line_comment = '#'
                multi_line_start = '=begin'
                multi_line_end = '=end'
            elif language in ['html', 'xml']:
                multi_line_start = '<!--'
                multi_line_end = '-->'
            
            # Extract single line comments
            single_line_comments = []
            if single_line_comment:
                pattern = f"{re.escape(single_line_comment)}(.+)$"
                matches = re.finditer(pattern, code_text, re.MULTILINE)
                for match in matches:
                    single_line_comments.append(match.group(1).strip())
            
            # Extract multi-line comments
            multi_line_comments = []
            if multi_line_start and multi_line_end:
                pattern = f"{re.escape(multi_line_start)}(.*?){re.escape(multi_line_end)}"
                matches = re.finditer(pattern, code_text, re.DOTALL)
                for match in matches:
                    multi_line_comments.append(match.group(1).strip())
            
            # Combine comments with code
            comments_text = "\n".join(single_line_comments + multi_line_comments)
            
            # Add metadata
            metadata['single_line_comments'] = len(single_line_comments)
            metadata['multi_line_comments'] = len(multi_line_comments)
            metadata['total_comments'] = len(single_line_comments) + len(multi_line_comments)
            metadata['code_lines'] = len(code_text.splitlines())
            
            # Return both code and extracted comments
            result_text = f"--- CODE ---\n{code_text}\n\n--- COMMENTS ---\n{comments_text}"
            
            logger.info(f"Successfully extracted code and comments: {metadata['total_comments']} comments from {metadata['code_lines']} lines")
            return result_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting code comments: {str(e)}")
            return f"[Code Extraction Error: {str(e)}]", {'error': str(e)}
    
    def extract_table_structure(self, pdf_content: bytes) -> List[str]:
        """
        Extract tables from PDF documents
        
        Args:
            pdf_content: Binary content of the PDF
            
        Returns:
            List of extracted tables as formatted strings
        """
        try:
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_content)
                temp_file_path = temp_file.name
            
            table_results = []
            
            # Try with tabula-py
            try:
                import tabula
                tables = tabula.read_pdf(temp_file_path, pages='all', multiple_tables=True)
                for i, table in enumerate(tables):
                    table_results.append(f"--- Table {i+1} ---\n{table.to_string()}")
            except Exception as tabula_error:
                logger.warning(f"Tabula extraction failed: {str(tabula_error)}")
                
                # Try with PyMuPDF's table extraction
                try:
                    doc = fitz.open(temp_file_path)
                    for page_num, page in enumerate(doc):
                        tables = page.find_tables()
                        for i, table in enumerate(tables):
                            if table.is_empty:
                                continue
                            rows = []
                            for row in table.rows:
                                row_data = []
                                for cell in row:
                                    text = page.get_text("text", cell)
                                    row_data.append(text.strip())
                                rows.append(" | ".join(row_data))
                            table_str = "\n".join(rows)
                            table_results.append(f"--- Table {page_num+1}-{i+1} ---\n{table_str}")
                except Exception as pymupdf_error:
                    logger.warning(f"PyMuPDF table extraction failed: {str(pymupdf_error)}")
            
            # Clean up temp file
            os.unlink(temp_file_path)
            
            logger.info(f"Successfully extracted {len(table_results)} tables from PDF")
            return table_results
            
        except Exception as e:
            logger.error(f"Error extracting tables from PDF: {str(e)}")
            return []
    
    def _detect_mime_type(self, file_content: bytes) -> str:
        """
        Detect MIME type from file content
        
        Args:
            file_content: Binary content of the file
            
        Returns:
            MIME type as string
        """
        try:
            import magic
            mime = magic.Magic(mime=True)
            return mime.from_buffer(file_content)
        except ImportError:
            # Fallback if python-magic is not installed
            logger.warning("python-magic not installed, using basic mime type detection")
            
            # Simple detection based on file signatures
            if file_content.startswith(b'%PDF'):
                return 'application/pdf'
            elif file_content.startswith(b'PK\x03\x04'):
                return 'application/zip'  # Could be DOCX, XLSX, etc.
            elif file_content.startswith(b'\xFF\xD8\xFF'):
                return 'image/jpeg'
            elif file_content.startswith(b'\x89PNG\r\n\x1A\n'):
                return 'image/png'
            else:
                return "application/octet-stream"
